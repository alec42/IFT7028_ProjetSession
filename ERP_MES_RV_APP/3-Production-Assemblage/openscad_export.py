import subprocess
import os
# import aspose.words as aw
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate#, Image
# import reportlab
from reportlab.lib.utils import ImageReader
from PIL import Image


def concatenate_bmp_images(image1_path, image2_path, output_path):
    # Open the first image
    image1_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), image1_path)
    image2_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), image2_path)
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), output_path)
    image1 = Image.open(image1_path)

    # Open the second image
    image2 = Image.open(image2_path)

    # Resize the images to have the same height
    height = min(image1.height, image2.height)
    image1 = image1.resize((int(image1.width * height / image1.height), height))
    image2 = image2.resize((int(image2.width * height / image2.height), height))

    # Calculate the total width including the separation
    separation = 20  # Width of the separation in pixels
    total_width = image1.width + separation + image2.width

    # Create a new image with the combined width and white background
    combined_image = Image.new("RGB", (total_width, height), (255, 255, 255))

    # Paste the first image on the left side
    combined_image.paste(image1, (0, 0))

    # Paste the second image on the right side with separation
    combined_image.paste(image2, (image1.width + separation, 0))

    # Save the combined image
    combined_image.save(output_path)
    
def add_projection_to_openscad(input_file_path, output_file_path):
    # Read the content of the input file
    with open(input_file_path, 'r') as file:
        content = file.read()
    
    # Split the content into translate blocks
    translate_blocks = content.split('translate(')[1:]

    # Generate the modified content with projection() added before each translate block
    modified_content = ''
    for block in translate_blocks:
        modified_content += 'projection() translate(' + block
    
    # Write the modified content to the output file
    with open(output_file_path, 'w') as file:
        file.write(modified_content)

def create_top_view_scad(original_scad_path, top_view_scad_path):
    with open(original_scad_path, "r") as original_file:
        original_script = original_file.read()

    # Split the original script into lines
    lines = original_script.split("\n")

    include_lines = []
    other_lines = []

    # Separate the use statements from the rest of the script
    for line in lines:
        if line.startswith("use"):
            include_lines.append(line)
        else:
            other_lines.append(line)
    new_line = "\n"
    top_view_script = f"""
    // Include statements from the original model
    {new_line.join(include_lines)}

    projection()  // Set the projection to top view, add translate([0,0,25]) rotate([90,0,0]) for side view

    // Rest of the original model
    {new_line.join(other_lines)}
    """

    with open(top_view_scad_path, "w") as file:
        file.write(top_view_script)

def export_to_bmp(input_path, output_path):
    project_root = os.path.dirname(os.path.abspath(__file__))
    # openscad_path = os.path.join(project_root, "OpenSCAD/openscad")
    openscad_path = os.path.join(project_root, "OpenSCAD/OpenSCAD.app/Contents/MacOS/OpenSCAD")
    output_svg = os.path.join(project_root, output_path)
    # top_view_scad_path = os.path.join(project_root, r"OpenSCAD_Parts/top_view.scad")
    top_view_scad_path = os.path.join(project_root, r"Test/top_view.scad")
    input_script = os.path.join(project_root, r"Test/top_view.scad")
    # input_script = os.path.join(project_root, r"OpenSCAD_Parts/top_view.scad")
    original_scad_path = os.path.join(project_root, input_path)

    # create_top_view_scad(original_scad_path, top_view_scad_path)
    
    
    if ".2d" in input_path:
        command = f'{openscad_path} -o {output_svg} {original_scad_path}'
    elif "beam" in input_path:
        add_projection_to_openscad(original_scad_path, top_view_scad_path)
        command = f'{openscad_path} -o {output_svg} {top_view_scad_path}'
    else:
        create_top_view_scad(original_scad_path, top_view_scad_path)
        command = f'{openscad_path} -o {output_svg} {top_view_scad_path}'
        
    # command = f'{openscad_path} -o {output_svg} {top_view_scad_path}'
    
    subprocess.call(command, shell=True)
    print("Finished openSCAD export. Creating png")
    svg_file = output_svg
    png_file = output_svg[:-3]+"bmp"
    # doc = aw.Document()
    # builder = aw.DocumentBuilder(doc)
    # shape = builder.insert_image(svg_file)
    # pageSetup = builder.page_setup
    # pageSetup.page_width = shape.width
    # pageSetup.page_height = shape.height
    # pageSetup.top_margin = 0
    # pageSetup.left_margin = 0
    # pageSetup.bottom_margin = 0
    # pageSetup.right_margin = 0
    # # save as PNG
    # doc.save(png_file)
    
    doc = Document()
    paragraph = doc.add_paragraph()
    picture = paragraph.add_run()
    picture.add_picture(png_file, width=Inches(6))
    doc.save(png_file)
    # doc = SimpleDocTemplate(png_file, pagesize=letter)
    # image = reportlab.platypus.Image(svg_file)#, width=letter[0], height=letter[1])
    # image.drawHeight = letter[0]
    # image.drawWidth  = letter[1]
    # story = [image]
    # doc.build(story)


def export_to_dxf(input_path, output_path):
    project_root = os.path.dirname(os.path.abspath(__file__))
    # openscad_path = os.path.join(project_root, "OpenSCAD/openscad")
    openscad_path = os.path.join(project_root, "OpenSCAD/OpenSCAD.app/Contents/MacOS/OpenSCAD")
    output_dxf = os.path.join(project_root, output_path)
    top_view_scad_path = os.path.join(project_root, r"Test/top_view.scad")
    input_script = os.path.join(project_root, r"Test/top_view.scad")
    # top_view_scad_path = os.path.join(project_root, r"OpenSCAD_Parts/top_view.scad")
    # input_script = os.path.join(project_root, r"OpenSCAD_Parts/top_view.scad")
    original_scad_path = os.path.join(project_root, input_path)
    if ".2d" in input_path:
        command = f'{openscad_path} -o {output_dxf} {input_path}'
    elif "beam" in input_path:
        add_projection_to_openscad(original_scad_path, top_view_scad_path)
        command = f'{openscad_path} -o {output_dxf} {top_view_scad_path}'
    else:
        create_top_view_scad(original_scad_path, top_view_scad_path)
        # print(f'{openscad_path} -o {output_dxf} top_view.scad')
        command = f'{openscad_path} -o {output_dxf} {top_view_scad_path}'


    subprocess.call(command, shell=True)


def generate_isometric_view(input_path, output_path):
    # Run OpenSCAD CLI command to generate the isometric view image
    project_root = os.path.dirname(os.path.abspath(__file__))
    output_png = os.path.join(project_root, output_path)
    input_scad = os.path.join(project_root, input_path)
    # openscad_path = os.path.join(project_root, "OpenSCAD/openscad")
    openscad_path = os.path.join(project_root, "OpenSCAD/OpenSCAD.app/Contents/MacOS/OpenSCAD")
    command = f'{openscad_path} -o {output_png} --camera=0,0,0,45,45,-45 --autocenter --viewall {input_scad}'
    subprocess.run(command, shell=True)
